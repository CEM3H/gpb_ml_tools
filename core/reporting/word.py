"""
Модуль для генерации отчетов в формате MS Word.
"""

from typing import Dict, List, Optional, Tuple, Any, Union
import os
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import io
from datetime import datetime

from ..container import ModelContainer
from ..metrics.classification import generate_confusion_matrix


class WordReportGenerator:
    """
    Класс для генерации отчетов о модели в формате MS Word.
    """
    
    def __init__(self, container: ModelContainer):
        """
        Инициализация генератора отчетов.
        
        Args:
            container: Контейнер с моделью и связанными артефактами
        """
        self.container = container
        self.document = Document()
        self._configure_document()
        
    def _configure_document(self):
        """Настройка документа - стили, шрифты и т.д."""
        # Задаем стиль заголовков
        style = self.document.styles['Title']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(16)
        font.bold = True
        
        # Настройка стиля Heading 1
        style = self.document.styles['Heading 1']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(14)
        font.bold = True
        
        # Настройка стиля Heading 2
        style = self.document.styles['Heading 2']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.bold = True
    
    def add_title(self, title: str):
        """Добавление заголовка отчета."""
        self.document.add_heading(title, level=0)
        # Добавляем дату создания отчета
        date_paragraph = self.document.add_paragraph()
        date_paragraph.add_run(f'Дата создания: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        
    def add_section(self, title: str, level: int = 1):
        """Добавление раздела отчета."""
        self.document.add_heading(title, level=level)
    
    def add_paragraph(self, text: str):
        """Добавление абзаца с текстом."""
        self.document.add_paragraph(text)
    
    def add_table(self, data: List[List[Any]], headers: List[str] = None, 
                 title: str = None, width: float = 6.0):
        """
        Добавление таблицы в отчет.
        
        Args:
            data: Данные для таблицы (список списков)
            headers: Заголовки столбцов
            title: Заголовок таблицы
            width: Ширина таблицы в дюймах
        """
        if title:
            self.add_section(title, level=2)
            
        rows = len(data) + (1 if headers else 0)
        cols = len(data[0]) if data else 0
        
        table = self.document.add_table(rows=rows, cols=cols)
        table.autofit = False
        table.width = Inches(width)
        
        # Добавляем заголовки, если они есть
        if headers:
            for i, header in enumerate(headers):
                cell = table.cell(0, i)
                cell.text = str(header)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Добавляем данные
        start_row = 1 if headers else 0
        for i, row_data in enumerate(data):
            for j, value in enumerate(row_data):
                cell = table.cell(i + start_row, j)
                cell.text = str(value)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавляем пустую строку после таблицы
        self.document.add_paragraph()
    
    def add_confusion_matrix(self, y_true: np.ndarray, y_pred: np.ndarray, 
                            title: str = "Матрица ошибок"):
        """
        Добавление матрицы ошибок в отчет.
        
        Args:
            y_true: Истинные значения
            y_pred: Предсказанные значения
            title: Заголовок раздела
        """
        self.add_section(title, level=2)
        
        # Генерируем матрицу ошибок
        cm, labels = generate_confusion_matrix(y_true, y_pred)
        
        # Преобразуем матрицу в список списков для таблицы
        headers = [''] + [f'Предсказанный класс {label}' for label in labels]
        data = []
        for i, row in enumerate(cm):
            data.append([f'Истинный класс {labels[i]}'] + list(row))
        
        # Добавляем таблицу с матрицей ошибок
        self.add_table(data, headers, title=None)
    
    def add_feature_importance(self, feature_importance: Dict[str, float], top_n: int = 10,
                              title: str = "Важность признаков"):
        """
        Добавление графика и таблицы с важностью признаков.
        
        Args:
            feature_importance: Словарь с важностью признаков
            top_n: Количество важнейших признаков для отображения
            title: Заголовок раздела
        """
        if not feature_importance:
            return
            
        self.add_section(title, level=2)
        
        # Сортируем признаки по важности
        sorted_features = sorted(feature_importance.items(), key=lambda x: x[1], reverse=True)
        if top_n:
            sorted_features = sorted_features[:top_n]
        
        # Создаем график
        plt.figure(figsize=(10, 6))
        feature_names = [f[0] for f in sorted_features]
        importances = [f[1] for f in sorted_features]
        
        # Строим горизонтальный столбчатый график
        plt.barh(range(len(feature_names)), importances, align='center')
        plt.yticks(range(len(feature_names)), feature_names)
        plt.xlabel('Важность')
        plt.title('Важность признаков')
        plt.tight_layout()
        
        # Сохраняем график в память
        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300)
        img_stream.seek(0)
        plt.close()
        
        # Добавляем изображение в документ
        self.document.add_picture(img_stream, width=Inches(6))
        self.document.add_paragraph()
        
        # Добавляем таблицу с важностью признаков
        headers = ['Признак', 'Важность']
        data = [[name, f"{importance:.6f}"] for name, importance in sorted_features]
        self.add_table(data, headers, title=None)
    
    def add_metrics_section(self, metrics: Dict[str, float], title: str = "Метрики модели"):
        """
        Добавление раздела с метриками модели.
        
        Args:
            metrics: Словарь с метриками модели
            title: Заголовок раздела
        """
        self.add_section(title, level=2)
        
        data = [[key, f"{value:.6f}" if isinstance(value, float) else str(value)] 
                for key, value in metrics.items()]
        headers = ['Метрика', 'Значение']
        self.add_table(data, headers, title=None)
    
    def add_model_info(self, title: str = "Информация о модели"):
        """
        Добавление информации о модели из метаданных.
        
        Args:
            title: Заголовок раздела
        """
        metadata = self.container.metadata
        self.add_section(title, level=1)
        
        # Основная информация
        data = [
            ['ID модели', metadata.model_id],
            ['Название модели', metadata.model_name],
            ['Автор', metadata.author],
            ['Дата создания', metadata.created_at],
            ['Описание целевой переменной', metadata.target_description]
        ]
        
        headers = ['Параметр', 'Значение']
        self.add_table(data, headers, title=None)
        
        # Таблицы, использованные для обучения
        if metadata.train_tables:
            train_tables_data = [[idx + 1, table] for idx, table in enumerate(metadata.train_tables)]
            self.add_table(train_tables_data, ['№', 'Таблица'], title="Таблицы для обучения")
    
    def generate_report(self, X_test: pd.DataFrame, y_test: pd.Series, 
                       output_path: str) -> str:
        """
        Генерация полного отчета о модели.
        
        Args:
            X_test: Тестовые данные (признаки)
            y_test: Тестовые данные (целевая переменная)
            output_path: Путь для сохранения отчета
            
        Returns:
            str: Путь к сохраненному отчету
        """
        self.add_title(f"Отчет о модели: {self.container.metadata.model_name}")
        self.add_model_info()
        
        # Добавляем секцию с метриками
        if self.container.metrics:
            self.add_metrics_section(self.container.metrics)
        
        # Получаем предсказания модели на тестовых данных
        if X_test is not None and y_test is not None:
            # Проверяем, может ли модель предсказывать вероятности
            has_predict_proba = hasattr(self.container.model, 'predict_proba')
            
            # Получаем предсказания
            y_pred = self.container.model.predict(X_test)
            
            # Для классификации добавляем матрицу ошибок
            if len(np.unique(y_test)) <= 10:  # Только для небольшого числа классов
                self.add_confusion_matrix(y_test, y_pred)
        
        # Добавляем важность признаков
        if self.container.feature_importance:
            self.add_feature_importance(self.container.feature_importance)
            
        # Проверяем наличие директории для отчета
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Сохраняем документ
        self.document.save(output_path)
        
        return output_path 